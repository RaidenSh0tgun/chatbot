"""
Batch test script for the SPAA chatbot.

Outputs:
- QA_test_with_chatbot_responses.xlsx
- QA_test_with_chatbot_responses.docx

Each Excel question receives a unique session_id, so test questions do not
share conversation history.

Required packages:
    pip install pandas openpyxl requests python-docx
"""

import re
import time
import uuid
from pathlib import Path
from typing import Any

import pandas as pd
import requests
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


# =========================
# Configuration
# =========================
INPUT_FILE = "QA_test.xlsx"
OUTPUT_EXCEL = "QA_test_with_chatbot_responses.xlsx"
OUTPUT_WORD = "QA_test_with_chatbot_responses.docx"

# Change this only if the route or port in main_revised_fast.py is different.
CHATBOT_URL = "http://127.0.0.1:5000/chat"

TIMEOUT_SECONDS = 180
DELAY_BETWEEN_REQUESTS = 0.5


# =========================
# API helpers
# =========================
def build_payload(question: str, session_id: str) -> dict[str, str]:
    """Build the JSON body expected by main_revised_fast.py."""
    return {
        "question": question,
        "session_id": session_id,
    }


def extract_response(data: Any) -> str:
    """Extract the chatbot text from common Flask JSON response structures."""
    if isinstance(data, str):
        return data

    if isinstance(data, dict):
        for key in ("response", "answer", "reply", "message", "text", "content"):
            if key in data and data[key] is not None:
                value = data[key]
                if isinstance(value, (dict, list)):
                    return str(value)
                return str(value)

        # Some APIs nest the answer inside a data/result object.
        for container_key in ("data", "result", "output"):
            if container_key in data:
                nested = extract_response(data[container_key])
                if nested:
                    return nested

    return str(data)


def find_question_column(df: pd.DataFrame) -> str:
    """Prefer a question-like column; otherwise use the first column."""
    normalized = {str(col).strip().lower(): col for col in df.columns}
    for candidate in ("question", "questions", "query", "prompt"):
        if candidate in normalized:
            return normalized[candidate]
    return df.columns[0]


def test_one_question(question: str, session_id: str) -> tuple[str, float, str, int | None]:
    """Send one independent question to the chatbot."""
    start_time = time.time()

    try:
        response = requests.post(
            CHATBOT_URL,
            json=build_payload(question, session_id),
            timeout=TIMEOUT_SECONDS,
        )
        elapsed = round(time.time() - start_time, 2)

        if not response.ok:
            # Preserve the backend's error message; this is critical for debugging 400s.
            body = response.text.strip()
            error = f"HTTP {response.status_code}: {body or response.reason}"
            return "", elapsed, error, response.status_code

        try:
            data = response.json()
        except ValueError:
            # Accept plain-text responses as well.
            return response.text.strip(), elapsed, "", response.status_code

        return extract_response(data).strip(), elapsed, "", response.status_code

    except requests.exceptions.Timeout:
        elapsed = round(time.time() - start_time, 2)
        return "", elapsed, f"Request timed out after {TIMEOUT_SECONDS} seconds", None
    except requests.exceptions.ConnectionError as exc:
        elapsed = round(time.time() - start_time, 2)
        return "", elapsed, f"Connection error: {exc}", None
    except requests.RequestException as exc:
        elapsed = round(time.time() - start_time, 2)
        return "", elapsed, f"Request error: {exc}", None
    except Exception as exc:
        elapsed = round(time.time() - start_time, 2)
        return "", elapsed, f"Unexpected error: {exc}", None


# =========================
# Word formatting helpers
# =========================
def add_inline_markdown(paragraph, text: str) -> None:
    """Render basic Markdown emphasis and inline code in a Word paragraph."""
    token_pattern = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)")
    position = 0

    for match in token_pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])

        token = match.group(0)
        if token.startswith(("**", "__")):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True

        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])


def add_markdown_to_document(document: Document, text: str) -> None:
    """Render common chatbot Markdown formatting into Word."""
    if not text:
        document.add_paragraph("[No response]")
        return

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    in_code_block = False
    code_lines: list[str] = []

    def flush_code_block() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["No Spacing"]
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            document.add_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) + 2, 9)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(paragraph, heading_match.group(2))
            continue

        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, bullet_match.group(1))
            continue

        number_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, number_match.group(1))
            continue

        quote_match = re.match(r"^>\s?(.*)$", stripped)
        if quote_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            run = paragraph.add_run(quote_match.group(1))
            run.italic = True
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, line)

    if in_code_block:
        flush_code_block()


def create_word_report(df: pd.DataFrame, question_col: str, output_path: Path) -> None:
    """Create a formatted Word report containing every question and response."""
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    document.add_heading("SPAA Chatbot Batch Test Results", level=0)
    document.add_paragraph(
        f"Total Excel rows: {len(df)}. Each valid question was tested with a unique session ID."
    )

    for position, (_, row) in enumerate(df.iterrows(), start=1):
        question = "" if pd.isna(row[question_col]) else str(row[question_col]).strip()
        answer = "" if pd.isna(row.get("chatbot_response", "")) else str(row.get("chatbot_response", ""))
        error = "" if pd.isna(row.get("error", "")) else str(row.get("error", ""))
        elapsed = row.get("response_time_seconds", "")
        status = row.get("http_status", "")

        document.add_heading(f"Test {position}", level=1)

        q_label = document.add_paragraph()
        q_run = q_label.add_run("Question")
        q_run.bold = True
        document.add_paragraph(question or "[Empty question]")

        r_label = document.add_paragraph()
        r_run = r_label.add_run("Chatbot response")
        r_run.bold = True
        add_markdown_to_document(document, answer)

        meta = document.add_paragraph()
        meta_run = meta.add_run("Test metadata: ")
        meta_run.bold = True
        meta.add_run(f"HTTP status={status}; response time={elapsed} seconds")

        if error:
            error_paragraph = document.add_paragraph()
            error_run = error_paragraph.add_run("Error: ")
            error_run.bold = True
            error_paragraph.add_run(error)

        if position < len(df):
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.save(output_path)


# =========================
# Main workflow
# =========================
def main() -> None:
    script_dir = Path(__file__).resolve().parent
    input_path = script_dir / INPUT_FILE
    excel_path = script_dir / OUTPUT_EXCEL
    word_path = script_dir / OUTPUT_WORD

    if not input_path.exists():
        raise FileNotFoundError(
            f"Cannot find {input_path}. Put QA_test.xlsx in the same folder as this script, "
            "or revise INPUT_FILE."
        )

    df = pd.read_excel(input_path)
    if df.empty:
        raise ValueError(f"{input_path.name} is empty.")

    question_col = find_question_column(df)

    responses: list[str] = []
    response_times: list[float] = []
    errors: list[str] = []
    status_codes: list[int | None] = []
    session_ids: list[str] = []

    total = len(df)

    for position, (_, row) in enumerate(df.iterrows(), start=1):
        raw_question = row[question_col]
        question = "" if pd.isna(raw_question) else str(raw_question).strip()
        session_id = f"batch_test_{uuid.uuid4().hex}"

        if not question:
            answer, elapsed, error, status = "", 0.0, "Empty question", None
            print(f"[{position}/{total}] Skipped empty question")
        else:
            print(f"[{position}/{total}] Testing: {question[:100]}")
            answer, elapsed, error, status = test_one_question(question, session_id)
            if error:
                print(f"  ERROR: {error}")
            else:
                print(f"  OK: {elapsed} seconds")
            time.sleep(DELAY_BETWEEN_REQUESTS)

        responses.append(answer)
        response_times.append(elapsed)
        errors.append(error)
        status_codes.append(status)
        session_ids.append(session_id)

    df["chatbot_response"] = responses
    df["response_time_seconds"] = response_times
    df["http_status"] = status_codes
    df["error"] = errors
    df["test_session_id"] = session_ids

    # Save both outputs independently so one export failure is clearly reported.
    df.to_excel(excel_path, index=False)
    print(f"\nExcel results saved to: {excel_path}")

    create_word_report(df, question_col, word_path)
    print(f"Word report saved to: {word_path}")


if __name__ == "__main__":
    main()
